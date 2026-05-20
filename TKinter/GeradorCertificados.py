import pandas as pd
from tkinter import *
from tkinter import messagebox
from tkinter import ttk

from docx import Document
from docx.shared import Pt

# tk - Biblioteca do tkinter
# Tk - Janela / Tela
janela = Tk()

# Insere um título para tela
janela.title("Gerador de Cerfificado")

# clam, alt, default, classic
stilo = ttk.Style()
stilo.theme_use("alt")
stilo.configure(".", font=("Arial 20"), rowheight=30)

treeviewDados = ttk.Treeview(janela, columns=(1, 2, 3, 4, 5, 6), show="headings")

treeviewDados.column("1", anchor=CENTER)
treeviewDados.heading("1", text="CPF")

treeviewDados.column("2", anchor=CENTER)
treeviewDados.heading("2", text="Nome")

treeviewDados.column("3", anchor=CENTER)
treeviewDados.heading("3", text="RG")

treeviewDados.column("4", anchor=CENTER)
treeviewDados.heading("4", text="Data Inicio")

treeviewDados.column("5", anchor=CENTER)
treeviewDados.heading("5", text="Data Fim")

treeviewDados.column("6", anchor=CENTER)
treeviewDados.heading("6", text="Email")

# grid - Divide a tela em grades / partes
# row - Linha
# column - Coluna
# columnspan - Colocamos para dizer quantas colunas do grid o item vai oculpar
# sticky - Usamos para preencher os espaços em brancos
# sticky - NSEW (Norte, Sul, Leste e Oeste)
treeviewDados.grid(row=4, column=0, columnspan=6, sticky="NSEW", pady=15)


def funcaoPassaDadosTreeviewParaEntry(Event):
    # Pego o 'item' selecionado
    item = treeviewDados.selection()

    # for - para
    for i in item:
        # Limpando os campos de entrada de dados
        exibirCPF.delete(0, END)
        exibirNome.delete(0, END)
        exibirRG.delete(0, END)
        exibirDataInicio.delete(0, END)
        exibirDataFim.delete(0, END)
        exibirEmail.delete(0, END)

        exibirCPF.insert(0, treeviewDados.item(i, "values")[0])
        exibirNome.insert(0, treeviewDados.item(i, "values")[1])
        exibirRG.insert(0, treeviewDados.item(i, "values")[2])
        exibirDataInicio.insert(0, treeviewDados.item(i, "values")[3])
        exibirDataFim.insert(0, treeviewDados.item(i, "values")[4])
        exibirEmail.insert(0, treeviewDados.item(i, "values")[5])


treeviewDados.bind("<Double-1>", funcaoPassaDadosTreeviewParaEntry)

# Abrindo o arquivo
dadosUsuarios = pd.read_excel("Dados.xlsx")
# dadosUsuarios = pd.read_excel("A:\\TKinter\\Certificados\\Dados.xlsx")

# print(dadosUsuarios)

# Convertendo a coluna Data Inicio para o tipo texto
dadosUsuarios["Data Inicio"] = dadosUsuarios["Data Inicio"].astype(str)

# Convertendo a coluna Data Fim para o tipo texto
dadosUsuarios["Data Fim"] = dadosUsuarios["Data Fim"].astype(str)

# for - para
for linha in range(len(dadosUsuarios)):
    # pegando o texto da coluna 3 que é a data
    # split - Quebra a coluna em partes de acordo com um critério
    # [0] - Pega na primeira parte
    dataInicioAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
    dataInicioMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
    dataInicioDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

    dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

    # --------------------------------------------

    dataFimAno = dadosUsuarios.iloc[linha, 4].split("-")[0]
    dataFimMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
    dataFimDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

    dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

    # Populando a Treeview com os dados do Excel
    treeviewDados.insert("", "end",
                         values=(str(dadosUsuarios.iloc[linha, 0]),  # CPF
                                 str(dadosUsuarios.iloc[linha, 1]),  # Nome
                                 str(dadosUsuarios.iloc[linha, 2]),  # RG
                                 str(dataInicioTratada),  # Data Inicio
                                 str(dataFimTratada),  # Data Fim
                                 str(dadosUsuarios.iloc[linha, 5])))  # Email

# grid   - Divide a tela em grades / partes
# row    - Linha
# column - Coluna
# sticky - Usamos para preencher os espaços em brancos
# sticky - NSEW (Norte, Sul, Leste e Oeste)
cpf = Label(text="CPF: ", font="Arial 12")
cpf.grid(row=0, column=0, sticky="E", pady=15)

exibirCPF = Entry(font="Arial 12")
exibirCPF.grid(row=0, column=1, sticky="W", pady=15)

# -------------------------------------------

nome = Label(text="Nome: ", font="Arial 12")
nome.grid(row=0, column=2, sticky="E", pady=15)

exibirNome = Entry(font="Arial 12")
exibirNome.grid(row=0, column=3, sticky="W", pady=15)

# -------------------------------------------

rg = Label(text="RG: ", font="Arial 12")
rg.grid(row=0, column=4, sticky="E", pady=15)

exibirRG = Entry(font="Arial 12")
exibirRG.grid(row=0, column=5, sticky="W", pady=15)

# -------------------------------------------

dataInicio = Label(text="Data Inicio: ", font="Arial 12")
dataInicio.grid(row=1, column=0, sticky="E", pady=15)

exibirDataInicio = Entry(font="Arial 12")
exibirDataInicio.grid(row=1, column=1, sticky="W", pady=15)

# -------------------------------------------

dataFim = Label(text="Data Fim: ", font="Arial 12")
dataFim.grid(row=1, column=2, sticky="E", pady=15)

exibirDataFim = Entry(font="Arial 12")
exibirDataFim.grid(row=1, column=3, sticky="W", pady=15)

# -------------------------------------------

email = Label(text="Email: ", font="Arial 12")
email.grid(row=1, column=4, sticky="E", pady=15)

exibirEmail = Entry(font="Arial 12")
exibirEmail.grid(row=1, column=5, sticky="W", pady=15)


def filtrarDados():
    print("Filtrando...")

    # for - para
    for linha in range(len(dadosUsuarios)):

        todasLinhas = treeviewDados.get_children()

        # Deletando todas as linhas da treeview
        treeviewDados.delete(*todasLinhas)

        # if - se
        # Se o CPF for vazio, não tiver nem uma informação eu carrego todos os dados
        if exibirCPF.get() == "":

            # Munda o texto do botão
            botaoPesquisar.config(text="Filtrar")

            # for - para
            for linha in range(len(dadosUsuarios)):
                # pegando o texto da coluna 3 que é a data
                # split - Quebra a coluna em partes de acordo com um criterio
                # [0] - Pegando primeira parte
                dataInicioAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
                dataInicioMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
                dataInicioDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

                dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

                # --------------------------------------------

                dataFimAno = dadosUsuarios.iloc[linha, 4].split("-")[0]
                dataFimMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
                dataFimDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

                dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

                # Populando a Treeview com os dados do Excel
                treeviewDados.insert("", "end",
                                     values=(str(dadosUsuarios.iloc[linha, 0]),  # CPF
                                             str(dadosUsuarios.iloc[linha, 1]),  # Nome
                                             str(dadosUsuarios.iloc[linha, 2]),  # RG
                                             str(dataInicioTratada),  # Data Inicio
                                             str(dataFimTratada),  # Data Fim
                                             str(dadosUsuarios.iloc[linha, 5])))  # Email

        else:

            # Munda o texto do botão
            botaoPesquisar.config(text="Limpar Filtros")

            # for - para
            for linha in range(len(dadosUsuarios)):

                # pegando o texto da coluna 3 que é a data
                # split - Quebra a coluna em partes de acordo com um criterio
                # [0] - Pegando primeira parte
                dataInicioAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
                dataInicioMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
                dataInicioDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

                dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

                # --------------------------------------------

                dataFimAno = dadosUsuarios.iloc[linha, 4].split("-")[0]
                dataFimMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
                dataFimDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

                dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

                # Verifico se o CPF do campo Entry é igual ao CPF da linha corrente do treeview
                if exibirCPF.get() == str(dadosUsuarios.iloc[linha, 0]):
                    # Populando a Treeview com os dados do Excel
                    treeviewDados.insert("", "end",
                                         values=(str(dadosUsuarios.iloc[linha, 0]),  # CPF
                                                 str(dadosUsuarios.iloc[linha, 1]),  # Nome
                                                 str(dadosUsuarios.iloc[linha, 2]),  # RG
                                                 str(dataInicioTratada),  # Data Inicio
                                                 str(dataFimTratada),  # Data Fim
                                                 str(dadosUsuarios.iloc[linha, 5])))  # Email


# grid       - Divide a tela em grades / partes
# row        - Linha
# column     - Coluna
# columnspan - Colocamos para dizer quantas colunas do grid o 'item' vai oculpar
# sticky     - Usamos para preencher os espaços em bracos
# sticky     - NSEW (Norte, Sul, Leste e Oeste)
botaoPesquisar = Button(text="PESQUISAR", font="Arial 14", command=filtrarDados)
botaoPesquisar.grid(row=5, column=0, columnspan=2, sticky="NSEW", padx=20)


def gerarCertificado():
    # Abrindo o arquivo do Word
    arquivoWord = Document("Certificado.docx")

    # Configurando os estilos
    estilo = arquivoWord.styles["Normal"]

    # Pegamos os dados do aluno dos campos Entry
    nomeAluno = exibirNome.get()
    dataInicio = exibirDataInicio.get()
    dataFim = exibirDataFim.get()
    nomeInstrutor = "Clevison Santos"
    CPF_Aluno = exibirCPF.get()
    RG_Aluno = exibirRG.get()

    frase_parte1 = " concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de "
    frase_montada = f"{nomeAluno}, CPF: {CPF_Aluno}, RG: {RG_Aluno}, {frase_parte1} {dataInicio} a {dataFim}."

    # for - para
    for paragrafo in arquivoWord.paragraphs:

        # if - se
        if "@nome" in paragrafo.text:
            paragrafo.text = nomeAluno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        if "@DataFim" in paragrafo.text:
            paragrafo.text = frase_montada
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    # Montando o caminho + nome do certificado
    # caminhoCertificadoGerado = "A:\\TKinter\\Certificados\\" + nomeAluno + ".docx"
    caminhoCertificadoGerado = nomeAluno + ".docx"

    # Salvando o certificado como o nome do aluno
    arquivoWord.save(caminhoCertificadoGerado)

    # Limpando os campos de entrada de dados
    exibirCPF.delete(0, END)
    exibirNome.delete(0, END)
    exibirRG.delete(0, END)
    exibirDataInicio.delete(0, END)
    exibirDataFim.delete(0, END)
    exibirEmail.delete(0, END)

    messagebox.showinfo("Mensagem", "Certificado gerado com sucesso!")


botaoGerarCertificado = Button(text="Gerar Certificado", font="Arial 14", command=gerarCertificado)
botaoGerarCertificado.grid(row=5, column=2, columnspan=2, sticky="NSEW", padx=20)


def gerarCertificadoEmMassa():
    # for - para
    # Passo linha por linha
    for linha in treeviewDados.get_children():

        # Os valores da linha corrente em forma de colunas
        coluna = treeviewDados.item(linha)["values"]

        # Pego as informações de cada coluna
        CPF_Separado = coluna[0]
        nomeAluno_Separado = coluna[1]
        rg_Separado = coluna[2]
        dataInicio_Separado = coluna[3]
        dataFim_Separado = coluna[4]
        nomeIntrutor_Separado = "Clevison Santos"

        # Abrindo o arquivo do Word
        arquivoWord = Document("Certificado.docx")

        # Configurando os estilos
        estilo = arquivoWord.styles["Normal"]

        frase_parte1 = " concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de "
        frase_montada = f"{nomeAluno_Separado}, CPF: {CPF_Separado}, RG: {rg_Separado}, {frase_parte1} {dataInicio_Separado} a {dataFim_Separado}."

        # for - para
        for paragrafo in arquivoWord.paragraphs:

            # if - se
            if "@nome" in paragrafo.text:
                paragrafo.text = nomeAluno_Separado
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

            if "@DataFim" in paragrafo.text:
                paragrafo.text = frase_montada
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

        # Montando o caminho + nome do certificado
        caminhoCertificadoGerado = "A:\\TKinter\\Certificados\\" + nomeAluno_Separado + ".docx"

        # Salvando o certificado como o nome do aluno
        arquivoWord.save(caminhoCertificadoGerado)

    messagebox.showinfo("Mensagem", "Certificados gerados com sucesso!")


botaoGerarCertificadoEmMassa = Button(text="Gerar em Massa", font="Arial 14", command=gerarCertificadoEmMassa)
botaoGerarCertificadoEmMassa.grid(row=5, column=4, columnspan=2, sticky="NSEW", padx=20)

# mainloop - No tkinter é uma janela em funcionamento como um loop
# A janela que o python mostra na verdade é um programa em funcionamento
janela.mainloop()
